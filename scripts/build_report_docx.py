#!/usr/bin/env python3
"""Build docs/Report02-apt-defense-lab.docx from docs/Google_Doc_Report_apt-defense-lab_UK.md"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    md_path = root / "docs" / "Google_Doc_Report_apt-defense-lab_UK.md"
    out_path = root / "docs" / "Report02-apt-defense-lab.docx"
    if not md_path.is_file():
        print(f"Missing {md_path}", file=sys.stderr)
        return 1

    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Pt(12)
        code_buf = []

    while i < len(lines):
        line = lines[i]
        raw = line

        if raw.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if raw.strip() == "---":
            i += 1
            continue

        # GitHub-style table
        if "|" in raw and raw.strip().startswith("|") and i + 1 < len(lines):
            sep = lines[i + 1].strip()
            if re.match(r"^\|?[\s\-:|]+\|?$", sep):
                headers = [c.strip() for c in raw.strip().split("|")[1:-1]]
                i += 2
                rows: list[list[str]] = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    row = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                    rows.append(row)
                    i += 1
                ncols = max(len(headers), max((len(r) for r in rows), default=0))
                table = doc.add_table(rows=1 + len(rows), cols=ncols)
                table.style = "Table Grid"
                for c, h in enumerate(headers[:ncols]):
                    table.rows[0].cells[c].text = h
                    _set_cell_shading(table.rows[0].cells[c], "E7E6E6")
                for r, row in enumerate(rows):
                    for c in range(ncols):
                        val = row[c] if c < len(row) else ""
                        table.rows[r + 1].cells[c].text = val
                doc.add_paragraph()
                continue

        if raw.startswith("# "):
            doc.add_heading(raw[2:].strip(), level=0)
        elif raw.startswith("## "):
            doc.add_heading(raw[3:].strip(), level=1)
        elif raw.startswith("### "):
            doc.add_heading(raw[4:].strip(), level=2)
        elif raw.startswith("- ") or raw.startswith("* "):
            p = doc.add_paragraph(raw[2:].strip(), style="List Bullet")
        elif raw.strip() == "":
            pass
        elif raw.startswith("**") and raw.rstrip().endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(raw.strip("*").strip())
            run.bold = True
        else:
            # strip minimal markdown bold inline
            t = raw
            t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
            p = doc.add_paragraph(t)
            for run in p.runs:
                if "**" in raw:
                    pass

        i += 1

    flush_code()
    doc.save(out_path)
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
